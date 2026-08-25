"""
Document parsing utilities.

Converts PDF / DOCX / TXT / Markdown / CSV files into plain text, with a
best-effort page/section split so that citations such as "Page 4" or
"Section 2" can be produced downstream.
"""
from __future__ import annotations

import csv
from dataclasses import dataclass, field
from pathlib import Path


@dataclass
class ParsedSection:
    """A logical section of a document (a page, or a paragraph block)."""

    label: str  # e.g. "Page 3" or "Section 2" or "Line 10"
    text: str


@dataclass
class ParsedDocument:
    document_name: str
    file_type: str
    full_text: str
    sections: list[ParsedSection] = field(default_factory=list)


class UnsupportedFileTypeError(Exception):
    pass


def parse_document(path: str | Path) -> ParsedDocument:
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix in (".txt", ".md"):
        return _parse_text(path)
    if suffix == ".csv":
        return _parse_csv(path)

    raise UnsupportedFileTypeError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> ParsedDocument:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    sections = []
    full_text_parts = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        sections.append(ParsedSection(label=f"Page {i}", text=text))
        full_text_parts.append(text)
    return ParsedDocument(
        document_name=path.name,
        file_type="pdf",
        full_text="\n".join(full_text_parts),
        sections=sections,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    import docx

    d = docx.Document(str(path))
    sections = []
    full_text_parts = []
    section_num = 0
    for para in d.paragraphs:
        if not para.text.strip():
            continue
        section_num += 1
        sections.append(ParsedSection(label=f"Section {section_num}", text=para.text))
        full_text_parts.append(para.text)
    return ParsedDocument(
        document_name=path.name,
        file_type="docx",
        full_text="\n".join(full_text_parts),
        sections=sections,
    )


def _parse_text(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="ignore")
    # Split on blank-line-separated blocks; label them as "Section N" if the
    # document uses "Section"/"Page" markers, else "Block N".
    raw_blocks = [b.strip() for b in text.split("\n\n") if b.strip()]
    sections = []
    for i, block in enumerate(raw_blocks, start=1):
        first_line = block.splitlines()[0]
        label = first_line if len(first_line) < 40 and any(
            k in first_line for k in ("Page", "Section", "Table", "Incident ID")
        ) else f"Block {i}"
        sections.append(ParsedSection(label=label, text=block))
    return ParsedDocument(
        document_name=path.name,
        file_type=path.suffix.lstrip("."),
        full_text=text,
        sections=sections,
    )


def _parse_csv(path: Path) -> ParsedDocument:
    rows_text = []
    sections = []
    with open(path, newline="", encoding="utf-8", errors="ignore") as f:
        # Skip leading '#' comment lines (used for the DEMO DATA label).
        lines = [line for line in f if not line.strip().startswith("#")]
    reader = csv.DictReader(lines)
    for i, row in enumerate(reader, start=1):
        row_text = ", ".join(f"{k}: {v}" for k, v in row.items())
        rows_text.append(row_text)
        sections.append(ParsedSection(label=f"Row {i}", text=row_text))
    return ParsedDocument(
        document_name=path.name,
        file_type="csv",
        full_text="\n".join(rows_text),
        sections=sections,
    )
